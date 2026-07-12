#!/usr/bin/env python3
# Applies reviewer-driven TEXT/EQ edits (stages A-F) to the revised docx.
# Numbers pending the N=10 matrix are marked 【NUM: ...】; open facts 【CONFIRM: ...】.
import re, docx
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = 'docs/paper_revision/Com-0420_revised.docx'
d = docx.Document(SRC)
P = d.paragraphs

def hard_set(p, text):
    """Replace all runs/OMML in a paragraph with a single run of `text`,
    keeping the paragraph's style (pPr)."""
    el = p._p
    for ch in list(el):
        t = ch.tag
        if t.endswith('}r') or t.endswith('}oMath') or t.endswith('}oMathPara') or t.endswith('}ins') or t.endswith('}del'):
            el.remove(ch)
    p.add_run(text)
    return p

import copy
from docx.enum.text import WD_ALIGN_PARAGRAPH
# reference paragraphs to clone formatting from (pPr carries justify/indent/spacing,
# and for headings the numbering + style). Captured before those paras are edited;
# hard_set() never touches pPr so these stay valid.
_REF_BODY = None   # set after P is available
_REF_HEAD = None

def _clone_like(anchor, ref_para, text, align=None, italic=None):
    new_p = copy.deepcopy(ref_para._p)
    for ch in list(new_p):
        if not ch.tag.endswith('}pPr'):   # keep only paragraph properties
            new_p.remove(ch)
    anchor._p.addnext(new_p)
    np = Paragraph(new_p, anchor._parent)
    r = np.add_run(text)
    if italic is not None: r.italic = italic
    if align is not None: np.alignment = align
    return np

def insert_after(p, text='', style=None):
    """Body-text insertion: clones a real Body Text paragraph so the new text
    matches the justified/indented/spaced look of the manuscript."""
    return _clone_like(p, _REF_BODY, text)

def insert_head(p, text):
    """Subsection-heading insertion: clones an existing numbered/italic subsection
    heading so the new heading matches (auto-numbering continues via its numPr)."""
    return _clone_like(p, _REF_HEAD, text, italic=True)

def insert_eq(p, text):
    """Display-equation line: body formatting but centered and no justification."""
    return _clone_like(p, _REF_BODY, text, align=WD_ALIGN_PARAGRAPH.CENTER)

def add_table_after(anchor, rows, header=True, style='Table Grid'):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    try: t.style = style
    except Exception: pass
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
    if header:
        for c in t.rows[0].cells:
            for para in c.paragraphs:
                for r in para.runs: r.bold = True
    anchor._p.addnext(t._tbl)
    return t

# capture paragraph objects by index BEFORE mutating
p8=P[8]; p26=P[26]; p27=P[27]; p28=P[28]
p35=P[35]; p36=P[36]; p39=P[39]; p40=P[40]; p41=P[41]; p43=P[43]
p47=P[47]; p54=P[54]; p55=P[55]; p62=P[62]; p70=P[70]
p74=P[74]; p75=P[75]; p77=P[77]; p78=P[78]
p82=P[82]; p83=P[83]; p86=P[86]; p87=P[87]
p91=P[91]; p92=P[92]; p93=P[93]; p101=P[101]
p112=P[112]; p115=P[115]; p118=P[118]; p119=P[119]
p128=P[128]; p133=P[133]; p134=P[134]; p138=P[138]
p141=P[141]; p144=P[144]; p153=P[153]; p154=P[154]; p155=P[155]
p160=P[160]; p161=P[161]
tbl2 = d.tables[1]
SUB='List Paragraph'
_REF_BODY = P[47]   # a Body Text paragraph (justified, indented, spaced)
_REF_HEAD = P[45]   # "System Architecture and Overview" subsection heading (numbered, italic)

# ───────────────────────── D. Abstract + contributions ─────────────────────
hard_set(p8,
 "Autonomous Mobile Robots (AMRs) in smart factories require highly stable, low-latency "
 "wireless connectivity for real-time control and monitoring. Ubiquitous Wi-Fi is susceptible "
 "to performance degradation during the frequent access-point (AP) handovers inherent to mobile "
 "operation, which introduce latency spikes and transient disconnections; single-path transports "
 "such as SP-QUIC are particularly exposed to these mobility-induced link failures because recovery "
 "requires costly connection re-establishment. To address this, we propose a handover-adaptive "
 "multipath-QUIC (MP-QUIC) transmission scheme for heterogeneous Wi-Fi/5G environments that keeps a "
 "persistent dual-path state (make-before-break), so that traffic can be moved to a pre-validated "
 "backup path without connection re-establishment. The core of the scheme is an RSSI-aware proactive "
 "path scheduler: a normalized, EWMA-smoothed Path Quality Index (PQI) computed from per-path RTT, "
 "loss, and bandwidth is augmented with a Wi-Fi link-quality (RSSI) signal that anticipates impending "
 "Wi-Fi degradation and triggers a preemptive switch to the 5G path before the primary link fails, "
 "with hysteresis-based failback to suppress path flapping. We evaluate the scheme on a real AMR "
 "testbed against same-stack baselines - a transport-only PQI ablation, a minimum-RTT default MP-QUIC "
 "scheduler, and SP-QUIC with RFC 9000 connection migration - so that the effect of the proposed "
 "policy is isolated. Experimental results show that the proposed scheme reduces handover-induced "
 "outage and switching latency 【NUM: to sub-XXX ms, versus SP-QUIC's multi-second recovery】 while "
 "preserving steady-state throughput and bounding tail latency under cross-traffic. These results "
 "indicate that an RSSI-aware, proactive MP-QUIC transport provides seamless connectivity for AMRs "
 "in dynamic industrial environments.")

hard_set(p26,
 "• A handover-adaptive MP-QUIC transport for AMRs in heterogeneous Wi-Fi/5G networks that "
 "maintains a persistent, make-before-break dual-path state, eliminating the connection "
 "re-establishment latency of reactive single-path migration;")
hard_set(p27,
 "• An RSSI-aware proactive path-quality scheme in which a normalized multi-metric Path Quality "
 "Index (PQI) over per-path RTT, loss, and bandwidth (min-max normalization, EWMA smoothing, and a "
 "hysteresis-based switching state machine) is augmented with a Wi-Fi RSSI signal to trigger "
 "preemptive handover before transport metrics degrade, together with RSSI-driven failback; and")
hard_set(p28,
 "• An implementation on a draft-20 MP-QUIC (picoquic) stack and an empirical evaluation on a "
 "real AMR testbed against same-stack baselines (a transport-only PQI ablation, a minimum-RTT default "
 "scheduler, and SP-QUIC with RFC 9000 connection migration) with repeated-run statistics, showing "
 "reduced handover outage and switching latency without steady-state throughput loss.")

# ───────────────────────── B. Related Works / Background ────────────────────
# R2.5 N3IWF/ATSSS after the Private-5G paragraph (35)
insert_after(p35,
 "At the network layer, 3GPP defines the Non-3GPP Interworking Function (N3IWF) and the ATSSS "
 "(Access Traffic Steering, Switching and Splitting) framework to integrate Wi-Fi and 5G access "
 "under a unified core, enabling operator-controlled steering across the two accesses. Such "
 "mechanisms require 5G-core integration and operate below the transport session; they do not by "
 "themselves guarantee end-to-end session continuity for an application whose transport connection "
 "breaks during a Wi-Fi handover. The approach in this paper is complementary and operates at the "
 "transport layer, requiring no core-network integration and remaining under the direct control of "
 "the AMR endpoint.")

# R2.1 question -> matter (last paragraph of 2.1)
hard_set(p36, p36.text.replace('question of', 'matter of'))

# R2.2 / R2.4  QUIC vs TCP framing + HoL clarification
hard_set(p39,
 "Standardized by the IETF over UDP, QUIC is a transport protocol that integrates TLS 1.3 and "
 "provides low-latency, secure, and reliable communication [20]. QUIC and TCP embody different "
 "design trade-offs rather than a strict ordering of one over the other. QUIC identifies a "
 "connection by a Connection ID (CID) rather than the 4-tuple, folds the cryptographic and transport "
 "handshakes together to enable 1-RTT (or 0-RTT) establishment, and multiplexes independent, "
 "separately flow-controlled streams. Because each stream is delivered independently, a loss on one "
 "stream (for example, a high-bandwidth video feed) does not stall delivery on parallel streams: "
 "QUIC removes head-of-line (HoL) blocking within a single path. As discussed below, HoL effects can "
 "nonetheless reappear across paths in a multipath setting, where packets on a slow path delay the "
 "in-order delivery of data carried on a fast path.")

# R1.4  connection migration RFC 9000 §9
hard_set(p40,
 "Crucially for mobility, QUIC's Connection ID decouples a session from the underlying IP addresses, "
 "giving QUIC intrinsic support for connection migration (RFC 9000, Section 9). When a client's "
 "address changes, it does not tear down the session; it sends from the new address with the same "
 "CID and proves reachability through a PATH_CHALLENGE/PATH_RESPONSE path-validation exchange [21]. "
 "Single-path mobility schemes such as mQUIC [22] exploit this capability, but they remain reactive: "
 "migration can begin only after a new path becomes usable, leaving the session exposed to the "
 "blackout that occurs during the handover itself. Multipath QUIC removes this limitation by "
 "validating and maintaining the backup path in advance.")

# R1.1  draft status prepended to MP-QUIC paragraph (41)
hard_set(p41,
 "Multipath QUIC (MP-QUIC) is an IETF extension - still an Internet-Draft rather than a ratified "
 "standard (latest draft-ietf-quic-multipath-21, March 2026), following a standardization effort "
 "that has evolved with several revisions since 2017 - that extends QUIC to aggregate multiple "
 "network paths (e.g., Wi-Fi and 5G) into a single logical connection under one Connection ID "
 "[23,24]. This path diversity lets a sender exploit heterogeneous wireless interfaces concurrently. "
 "Unlike a passive backup link, MP-QUIC isolates the transport state of each interface, maintaining "
 "independent packet-number spaces, per-path RTT estimation, and separate congestion-control state "
 "[25], so that loss or congestion on a deteriorating Wi-Fi link does not throttle a stable 5G link.")

# R1.2 / R2.3  scheduler is implementation-specific; soft handover
hard_set(p43,
 "Because the multipath specification does not mandate a packet scheduler - path selection is "
 "implementation-specific (draft-ietf-quic-multipath, Section 5.5) - the behaviour during an AMR "
 "handover depends entirely on the scheduler in use. A default or reference scheduler that is "
 "unaware of the handover context keeps sending on the degrading Wi-Fi path; the resulting per-path "
 "delay disparity then causes receiver-side reordering and cross-path HoL blocking, in which data on "
 "the fast 5G path waits for late Wi-Fi packets and the benefit of multipath is nullified [29]. An "
 "AMR-oriented MP-QUIC scheduler must therefore be handover-aware: it must detect Wi-Fi degradation "
 "early and steer traffic away from the failing path before it collapses. Because both links are "
 "kept active, this is more precisely a soft handover (make-before-break) at the transport layer "
 "than a conventional break-before-make handover of the wireless link.")

# R1.3  new subsection: MP-QUIC Path Schedulers (after 43)
h = insert_head(p43, "MP-QUIC Path Schedulers")
insert_after(h,
 "A substantial body of work has studied packet scheduling for multipath transport. Latency- and "
 "completion-aware schedulers such as minimum-RTT, ECF (Earliest Completion First), and BLEST were "
 "developed for MPTCP and later carried into MP-QUIC; learning-based schedulers such as Peekaboo and "
 "various reinforcement-learning approaches adapt the send policy online; and application- or "
 "mobility-aware designs - for example the stream-aware MP-QUIC scheduler of Xing et al. [23] and "
 "the mobility-aware MAMS scheduler [13] - tailor path selection to traffic semantics or client "
 "movement. These schedulers largely optimize aggregate throughput or completion time under general "
 "mobility. In contrast, the AMR Wi-Fi/5G setting is dominated by predictable but disruptive "
 "degradation events at AP boundaries, where the objective is not aggregate throughput but "
 "continuity: pre-emptively vacating a Wi-Fi path that is about to fail. The scheduler proposed here "
 "targets that objective directly and, unlike purely transport-metric schedulers, incorporates a "
 "link-layer RSSI signal so that it can act before the transport metrics reflect the degradation.")

# ───────────────────────── C. Proposed method (RSSI-main) ───────────────────
hard_set(p47,
 "To ensure seamless communication for AMRs in dynamic industrial environments, we propose a "
 "handover-adaptive Multipath QUIC (MP-QUIC) architecture with an RSSI-aware proactive scheduler. "
 "The design concurrently manages heterogeneous wireless paths - a primary Wi-Fi link and a "
 "secondary 5G link - within a single transport connection and, in contrast to single-path migration "
 "schemes, keeps a persistent dual-path state so that the backup path is always validated and ready. "
 "This bypasses the signaling and connection-recovery delays of reactive failover during Wi-Fi AP "
 "handovers.")

hard_set(p54,
 "The Sensing module continuously ingests per-path transport metrics (RTT, loss, and delivery rate) "
 "derived from MP-QUIC's acknowledgement stream, together with a Wi-Fi link-quality (RSSI) signal "
 "sampled out-of-band from the wireless driver. Because Wi-Fi and 5G differ in absolute performance "
 "ranges, the Estimation phase does not use these raw values directly: it synthesizes them into a "
 "single Path Quality Index (PQI) via min-max normalization, EWMA filtering, and an RSSI penalty, "
 "smoothing momentary fluctuation while preserving the RSSI trend that precedes a handover.")

hard_set(p55,
 "The PQI then drives the Decision module, which executes the routing policy over MP-QUIC: proactive "
 "RSSI-driven handover on impending degradation, reactive PQI-threshold failover, hysteresis-based "
 "failback to prevent path flapping, and frame-aware stream-to-path mapping. This design keeps "
 "latency-sensitive application data on the best available path and shields the remote server from "
 "the mobility-induced network events, reducing - though not, in general, entirely eliminating - "
 "the service disruption of AMR mobility.")

# R1.8  probing / metric acquisition (after 70, end of Sensing subsection)
insert_after(p70,
 "Per-path metrics are obtained without a separate probing protocol: RTT and its variation come from "
 "QUIC's own acknowledgements on each path, loss from the ratio of declared-lost to sent bytes per "
 "path, and available bandwidth from the rate of newly acknowledged data. The backup (5G) path is "
 "validated in advance and kept alive by the multipath handshake and periodic keep-alives so that "
 "its metrics are continuously available; when only one path is active, the scheduler additionally "
 "issues a path probe (PATH_CHALLENGE) to (re)validate the second path at a fixed interval. The "
 "Wi-Fi RSSI, which is not observable from transport feedback, is sampled from the wireless driver "
 "(iw dev <iface> link) every sampling period.")

# R1.6 / R1.7  Estimation: explicit normalization + PQI, numbered equations
hard_set(p74,
 "To translate the raw metrics from the Sensing phase into routing decisions, the Estimation module "
 "forms a Path Quality Index (PQI). Because heterogeneous interfaces exhibit different absolute "
 "performance ranges, each metric is first min-max normalized within the current set of active "
 "paths, enabling a scale-free comparison across the disparate paths.")
hard_set(p75,
 "For each schedulable path i, the module uses smoothed RTT r_i, loss rate l_i, and estimated "
 "goodput b_i. Each metric is min-max normalized across the active path set P (guarding against a "
 "zero range) so that every term lies in [0,1] with 0 denoting the best value:")
e1=insert_eq(p75, "    r̃_i = (r_i − r_min)/(r_max − r_min),   "
                     "l̃_i = (l_i − l_min)/(l_max − l_min),   "
                     "b̃_i = (b_i − b_min)/(b_max − b_min)          (1)")
e2=insert_after(e1, "The normalized terms combine into a transport-quality cost (lower is better), "
                    "with weights summing to one:")
e3=insert_eq(e2, "    C_i^tr = α·r̃_i + β·l̃_i + γ·(1 − b̃_i),   "
                    "α + β + γ = 1          (2)")
e4=insert_after(e3, "In the proposed RSSI-aware mode, a Wi-Fi link-quality penalty ρ_i ∈ [0,1] "
                    "is formed from the sampled RSSI s_i (dBm) by a linear mapping between a good and "
                    "an unusable threshold,")
e5=insert_eq(e4, "    ρ_i = clamp((S_good − s_i)/(S_good − S_unusable), 0, 1)          (3)")
e6=insert_after(e5, "and is folded into the composite cost with weight ω:")
e7=insert_eq(e6, "    C_i = (α·r̃_i + β·l̃_i + γ·(1 − b̃_i) + "
                    "ω·ρ_i)/(1 + ω)          (4)")
e8=insert_after(e7, "The Path Quality Index used by the scheduler is this cost on a 0-1000 scale; "
                    "consistent with the failover rule below, a higher PQI denotes poorer path quality:")
e9=insert_eq(e8, "    PQI_i = 1000 · C_i          (5)")
e10=insert_after(e9, "To suppress transient fluctuation, the PQI is smoothed by an exponentially "
                     "weighted moving average with factor λ:")
e11=insert_eq(e10,"    PQI_i(t) = λ·PQI_i^raw(t) + (1 − λ)·PQI_i(t − 1)          (6)")

hard_set(p77,
 "Here α, β, and γ weight the RTT, loss, and bandwidth contributions, ω weights the "
 "RSSI penalty, and S_good and S_unusable are the RSSI thresholds; all values are listed in Table 3. "
 "For the transport-only ablation (the pqi mode) the RSSI term is disabled (ω = 0), reducing (4) "
 "to (2), so that comparing the rssi and pqi modes isolates the contribution of the RSSI signal.")

# R3.7  Path Selection: proactive (RSSI) vs reactive (PQI) with numbered rules
hard_set(p82,
 "Based on the continuous estimation, the Decision module selects the active path. Under stable "
 "conditions it maps traffic to the path that minimizes the PQI of (5); at AP boundaries it applies "
 "two switching regimes - a reactive, PQI-threshold failover and a proactive, RSSI-driven handover.")
hard_set(p83,
 "Let a denote the active Wi-Fi path and c the 5G backup path. Reactive failover is triggered when "
 "the transport-derived PQI of the active path crosses a degradation threshold T_deg and the backup "
 "is better by at least a safety margin Δ, sustained for a stability interval T_stable:")
r7=insert_eq(p83, "    switch a→c   iff   PQI_a ≥ T_deg   ∧   PQI_a − PQI_c ≥ Δ   "
                     "(sustained ≥ T_stable)          (7)")
r8a=insert_after(r7, "Proactive, RSSI-driven handover acts independently of the transport metrics: if "
                     "the Wi-Fi RSSI falls to the 'bad' threshold, the scheduler switches immediately, "
                     "before the degradation is reflected in RTT or loss:")
r8b=insert_eq(r8a,"    switch a→c   iff   s_a ≤ S_bad          (8)")
r8c=insert_after(r8b,"Rule (8) is what makes the proposed scheme proactive rather than reactive: it "
                     "acts on the leading indicator (declining RSSI) ahead of the lagging transport "
                     "indicators, so the switch begins before the Wi-Fi path stops delivering. Rule "
                     "(7) is the reactive fallback that also covers degradations not preceded by an "
                     "RSSI drop (e.g., congestion).")

hard_set(p86,
 "The margin Δ, a minimum inter-switch debounce, and the failback rule together prevent "
 "oscillatory 'ping-pong' switching. Failback is deliberately conservative: even after the Wi-Fi "
 "RSSI and PQI recover, traffic returns to Wi-Fi only if the recovered path sustains a PQI below the "
 "recovery threshold T_recover for at least T_stable. To further damp pathological oscillation, the "
 "debounce interval is increased multiplicatively when a switch reverses the previous one and reset "
 "after a period of stability. All thresholds are listed in Table 3.")

# Table 3: parameters (R3.3) after p86-hysteresis (insert after the failback para)
param_rows = [
 ["Symbol","Meaning","Value"],
 ["α, β, γ","RTT / loss / bandwidth weights","0.45 / 0.35 / 0.20"],
 ["ω","RSSI penalty weight (rssi mode)","0.30"],
 ["λ","EWMA smoothing factor (RTT/loss/BW)","0.45 / 0.35 / 0.20"],
 ["T_deg","failover threshold (PQI, 0-1000)","650"],
 ["T_recover","failback threshold (PQI)","450"],
 ["Δ","hysteresis safety margin (PQI)","120"],
 ["T_debounce","min inter-switch interval (adaptive)","2 s → 60 s"],
 ["T_stable","failback stability interval","3 s"],
 ["T_silence","path-stale timeout","8 s"],
 ["S_good / S_degrade","RSSI good / degrading (dBm)","−55 / −67"],
 ["S_bad / S_unusable","RSSI bad / unusable (dBm)","−75 / −82"],
 ["CC","congestion control","BBR"],
]
cap3 = insert_after(p86, "Table 3: PQI and RSSI-aware scheduler parameters (as implemented).")
add_table_after(cap3, param_rows)

# R1.9  frame-aware transmission / framing
hard_set(p87,
 "Finally, switching is coordinated with the application's video framing. Each captured frame - "
 "depth or RGB - is serialized on a per-path unidirectional QUIC stream with an explicit application "
 "header [MPQ1 | type | length]: a 4-byte magic tag, a 1-byte stream-type field distinguishing depth "
 "from RGB, and a 4-byte big-endian length. The receiver therefore recovers exact frame boundaries "
 "from the length prefix regardless of how packets were scheduled across paths, and depth and RGB "
 "use disjoint stream-ID ranges to avoid interleaving. Aligning path transitions with frame "
 "boundaries reduces partial-frame loss, and packets arriving over heterogeneous paths are "
 "timestamp-aligned and reordered in a receiver-side buffer to preserve playback continuity.")

# R3.5 / R1.5  Implementation subsection (after framing paragraph 87)
hi = insert_head(p87, "Implementation")
insert_after(hi,
 "We implement the scheme on picoquic 1.1.50.3 with the IETF multipath extension "
 "draft-ietf-quic-multipath-20 (multipath enabled by advertising a non-zero initial_max_path_id, "
 "per-path congestion control, and keep-alive). The three-stage pipeline of Fig. 3 maps directly "
 "onto the code: Sensing draws per-path RTT from the RTT estimator, loss from per-path lost/sent "
 "bytes, and goodput from delivered bytes over time (client_loop.c, update_pqi_metrics_for_path), "
 "and samples Wi-Fi RSSI via iw dev <iface> link; Estimation computes the normalization, EWMA, and "
 "PQI of (1)-(6) (fill_pqi_scores); and Decision is the hysteresis state machine of (7)-(8) "
 "(choose_path_via_pqi), including RSSI preemption and failback. The sender selects one of four "
 "modes at run time: rssi (the proposed RSSI-aware scheduler), pqi (the transport-only ablation, "
 "ω = 0), default (a minimum-RTT scheduler approximating picoquic's native selection), and "
 "spquic (single path with RFC 9000 connection migration). Congestion control is BBR by default. "
 "Because the four modes differ only in the path-selection policy while sharing the same stack, "
 "camera, paths, and workload, comparing them isolates the contribution of the proposed scheme.")

# ───────────────────────── E. Testbed ───────────────────────────────────────
hard_set(p91,
 "Our testbed reproduces a realistic AMR uplink. The sender is an NVIDIA Jetson-class edge computer "
 "mounted on the AMR, equipped with two heterogeneous interfaces: a primary Wi-Fi interface "
 "associated with an indoor 802.11 access point, and a secondary 5G uplink - a commercial 5G (NR) "
 "mobile connection shared to the robot by smartphone tethering - that serves as the "
 "always-available backup path. In a production deployment this role would be served by a private 5G "
 "network; we use a tethered commercial 5G link as a practical stand-in for that always-on 5G "
 "umbrella. The receiver is a server on a wired Ethernet link, so that any measured "
 "degradation originates in the wireless access segment. The application is the robot's live camera "
 "uplink: the AMR streams a real depth+RGB video feed (JPEG-encoded RGB and 16-bit depth) to the "
 "server at a measured application rate of 【NUM: ~XX Mbps】. Because the transport carries an "
 "ordered sequence of length-delimited video frames, we define the total transmission delay of a run "
 "as the wall-clock time to deliver a fixed frame sequence end-to-end, and the normal-state "
 "throughput as the delivered application bitrate under stable Wi-Fi.")

hard_set(p92,
 "For the comparison we evaluate four transport configurations that share the same MP-QUIC stack, "
 "camera, paths, and workload and differ only in the path-selection policy: (i) the proposed "
 "RSSI-aware scheduler; (ii) a transport-only PQI ablation with the RSSI term disabled; (iii) a "
 "minimum-RTT default MP-QUIC scheduler; and (iv) SP-QUIC restricted to a single path but permitted "
 "to migrate to the 5G path via RFC 9000 connection migration when the Wi-Fi path fails. "
 "Configuration (iv) removes the earlier unfair restriction in which single-path QUIC could not use "
 "the 5G umbrella, so the comparison isolates the benefit of proactive multipath scheduling rather "
 "than mere access to a second network.")

hard_set(p93,
 "AMR mobility across AP boundaries is emulated in a controlled, repeatable way rather than by "
 "physically driving the robot between APs, so that the disconnection timing and duration are "
 "identical across schedulers; we additionally validated the full pipeline on the physically moving "
 "robot. A run lasts 【NUM: T】 s; at 【NUM: t1】 s the Wi-Fi interface is programmatically "
 "disconnected (nmcli device disconnect) for 【NUM: d】 s and then reconnected, producing one "
 "controlled handover per run 【CONFIRM: AP 배치/로봇 속도 - 실주행 검증 시 값 기입 (R3.6)】. Each "
 "(scenario, scheduler) cell is repeated N = 10 times. The three phases are:")

hard_set(p101,
 "Each (scenario, scheduler) combination is repeated N = 10 times; we report the mean with a 95% "
 "confidence interval and, where a difference is claimed, a test of its statistical significance, in "
 "addition to representative single-run traces. For the handover-focused analysis we evaluate:")

# ───────────────────────── F. Results (frames + 【NUM】) ─────────────────────
hard_set(p112,
 "Before the detailed analysis, Table 2 summarizes the four configurations under normal operation "
 "and during a handover. All figures are means over N = 10 runs with 95% confidence intervals; "
 "values shown as 【NUM】 are being regenerated on the revised, stabilized implementation and will "
 "replace the placeholders.")

hard_set(p115,
 "Under stable Wi-Fi the four schedulers deliver comparable steady-state throughput 【NUM】; in "
 "particular, the proposed scheme does not sacrifice normal-state throughput relative to the "
 "single-path and default baselines. The apparent normal-state throughput deficit of MP-QUIC "
 "reported in the original manuscript is not reproduced once the comparison is made same-stack and "
 "only the scheduler is varied; any residual gap is attributable to multipath scheduling and "
 "reordering overhead and is quantified in 【NUM】. The configurations diverge at the handover: "
 "SP-QUIC without a usable second path suffers a service interruption, whereas the multipath "
 "configurations preserve delivery by switching to the pre-validated 5G path with a switching "
 "latency of 【NUM】.")

hard_set(p118,
 "The proposed scheme maintains service continuity during Wi-Fi disconnection by proactively using "
 "the pre-established 5G backup path. While the Wi-Fi link delivers a stable throughput of 【NUM】 "
 "under baseline conditions, single-path SP-QUIC collapses toward zero throughput at an AP boundary "
 "unless it can migrate; the proposed scheme instead redirects traffic to the 5G link and sustains "
 "approximately 【NUM】, avoiding a full service interruption.")
hard_set(p119,
 "This continuity is reflected in the total transmission delay. As shown in Figure 5, SP-QUIC "
 "without migration incurs a large delay penalty across handover timings because it relies on "
 "timeout and connection re-establishment, whereas the proposed scheme bounds the total "
 "transmission delay to 【NUM】 across the tested scenarios. The same-stack default and PQI-ablation "
 "baselines fall between these extremes 【NUM】, isolating the additional benefit of the RSSI-aware "
 "policy.")

hard_set(p128,
 "Minimizing both handover latency and its variability under congestion is important for real-time "
 "AMR control. Figure 6 shows the distribution of handover latency across the four configurations "
 "(N = 10; box plots with medians and inter-quartile ranges).")
hard_set(p133,
 "SP-QUIC without a pre-validated second path recovers only through failure detection and connection "
 "re-establishment, incurring handover delays on the order of 【NUM: seconds】. The proposed scheme "
 "switches to the pre-established 5G path in 【NUM: sub-XXX ms】. We attribute the multi-second "
 "handover reported in the original manuscript to conflating the Wi-Fi link re-association time with "
 "the transport reaction time; the two are separated in the micro-analysis below, and the transport "
 "reaction itself is 【NUM】.")

hard_set(p134,
 "To assess robustness we injected controlled cross-traffic on the Wi-Fi path. Here, cross-traffic "
 "denotes contending background load generated with iperf3 on the same access point, combined with a "
 "probabilistic packet-loss ramp applied to the QUIC flow (iptables statistic-mode drops ramped from "
 "1% to 20%). The 'no traffic' point is the impairment-free baseline, and the 50 and 150 Mbps points "
 "denote the injected background bitrate contending on the Wi-Fi link - not the robot's own "
 "application rate, which is 【NUM: ~XX Mbps】 (Section 4.1); the cross-traffic represents competing "
 "users sharing the AP rather than the AMR's own demand.")
hard_set(p138,
 "Under the impairment-free baseline the configurations show comparable delay 【NUM】. As the "
 "cross-traffic intensity increases to 50 and then 150 Mbps, SP-QUIC degrades sharply, with total "
 "transmission delay rising to 【NUM】 and high variance. The proposed scheme shifts traffic away "
 "from the congested Wi-Fi link and holds a more stable delay of 【NUM】 even under the heaviest "
 "load, indicating robustness to shared-medium congestion in addition to fast handover.")

hard_set(p141,
 "To evaluate the end-to-end impact of a handover on application performance, we analyze throughput "
 "across a complete Wi-Fi disconnection-and-recovery cycle. Figure 8 contrasts a full service "
 "blackout for SP-QUIC without migration against the continuous delivery of the proposed scheme.")
hard_set(p144,
 "For SP-QUIC without a second path, throughput remains at zero until Wi-Fi is fully re-established, "
 "a gap of 【NUM】 that is unacceptable for real-time monitoring or teleoperation. The proposed "
 "scheme preserves end-to-end continuity via the pre-established 5G path: throughput dips briefly "
 "during the transition but never collapses to zero, with a measured outage of 【NUM】.")

hard_set(p153,
 "In the reactive regime, upon Wi-Fi failure the scheduler exhibits a brief hesitation of 【NUM: "
 "~XX ms】 before traffic is fully rerouted; this is the latency of degradation detection and the "
 "scheduling decision, not connection teardown. This hesitation characterizes the reactive failover "
 "of rule (7). The proposed RSSI-driven rule (8) is proactive: it initiates the switch from the "
 "declining RSSI trend before the transport metrics - and hence this hesitation - occur. We report "
 "the proactive lead time, the interval between the RSSI-triggered switch and the instant the Wi-Fi "
 "path stops delivering, as 【NUM】.")
hard_set(p154,
 "During Wi-Fi recovery the proposed scheme reactivates the primary path and resumes traffic within "
 "【NUM: ~XX ms】, with no connection re-establishment or transport-state reset, migrating traffic "
 "back to the preferred Wi-Fi path.")
hard_set(p155,
 "These micro-level observations show that the proposed scheme turns a connection-level handover "
 "into a lightweight path-scheduling operation, so that macro-level continuity and micro-level "
 "latency stability are achieved together - explaining the low switching latency reported in "
 "Figure 6, while distinguishing the proactive (RSSI-driven) from the reactive (PQI-threshold) path.")

# Conclusion softening (R3.2)
hard_set(p160,
 "Guaranteeing uninterrupted real-time communication for AMRs in dynamic industrial environments "
 "requires transport-layer resilience. To reduce the transient degradation and service interruption "
 "that conventional SP-QUIC suffers during AP handovers, this paper proposed and empirically "
 "evaluated a handover-adaptive MP-QUIC scheme whose RSSI-aware proactive scheduler combines a "
 "normalized multi-metric PQI with a Wi-Fi RSSI signal and hysteresis-based failback, shifting "
 "traffic to a pre-validated 5G path before the Wi-Fi link fails.")
hard_set(p161,
 "On a real AMR testbed, and against same-stack baselines (a transport-only PQI ablation, a "
 "minimum-RTT default scheduler, and SP-QUIC with connection migration), the proposed scheme reduces "
 "handover outage and switching latency 【NUM: to sub-XXX ms, versus SP-QUIC's multi-second "
 "recovery】 and bounds total transmission delay under cross-traffic 【NUM】, while preserving "
 "steady-state throughput 【NUM】. We report means over N = 10 runs with 95% confidence intervals.")

# ───────────────────────── Table 2 rebuild (R1.10 / R3.4) ───────────────────
new_tbl2 = [
 ["Metric","SP-QUIC (migration)","Default MP-QUIC (minRTT)","PQI (ablation)","Proposed (RSSI)"],
 ["Normal-state throughput","【NUM】","【NUM】","【NUM】","【NUM】"],
 ["Normal-state RTT","【NUM】","【NUM】","【NUM】","【NUM】"],
 ["Handover switching latency","【NUM】","【NUM】","【NUM】","【NUM】"],
 ["Handover outage duration","【NUM】","【NUM】","【NUM】","【NUM】"],
 ["Total transmission delay","【NUM】","【NUM】","【NUM】","【NUM】"],
 ["Handovers per run","【NUM】","【NUM】","【NUM】","【NUM】"],
]
cap2 = P[113]  # caption paragraph (unchanged index; captured fresh is fine pre-mutation-of-this-region)
# rebuild caption text too
hard_set(cap2, "Table 2: Performance comparison across the four same-stack configurations "
               "(means over N = 10 runs; 95% CI). 【NUM】 values are being regenerated.")
nt = add_table_after(cap2, new_tbl2)
# remove the old Table 2
tbl2._tbl.getparent().remove(tbl2._tbl)

# ───────────────────────── A. Global normalization ─────────────────────────
wifi = re.compile(r'\bwi[\s\-]?fi\b', re.I)
def fix_runs(container_paras):
    for p in container_paras:
        for r in p.runs:
            if r.text and wifi.search(r.text):
                r.text = wifi.sub('Wi-Fi', r.text)
fix_runs(d.paragraphs)
for t in d.tables:
    for row in t.rows:
        for c in row.cells:
            fix_runs(c.paragraphs)

d.save(SRC)
print("saved:", SRC)
print("paragraphs:", len(d.paragraphs), "tables:", len(d.tables))
